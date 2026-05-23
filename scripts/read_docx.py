"""Extract text from Word documents (.docx) for Claude to read."""
import sys
import os

try:
    import docx
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def extract_docx(filepath):
    """Extract all text from a .docx file, preserving paragraph structure."""
    doc = docx.Document(filepath)

    parts = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            # Show heading style if present
            if para.style.name.startswith("Heading"):
                parts.append(f"\n## {para.text}")
            else:
                parts.append(para.text)

    # Tables
    for table in doc.tables:
        parts.append("\n--- TABLE ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python scripts/read_docx.py <file.docx> [--raw]")
        sys.exit(1)

    path = sys.argv[1]
    if not os.path.exists(path):
        print(f"Error: file not found: {path}")
        sys.exit(1)

    if not path.lower().endswith(".docx"):
        print(f"Error: only .docx files are supported. For .doc, convert with LibreOffice first.")
        sys.exit(1)

    text = extract_docx(path)
    # Wrap stdout to handle GBK-incompatible characters on Windows
    sys.stdout.reconfigure(encoding="utf-8", errors="replace") if hasattr(sys.stdout, "reconfigure") else None
    print(text, flush=True)
